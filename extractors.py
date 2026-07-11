"""
طبقة الاستخراج الموحّدة — تحوّل أي ملف مدخل إلى جدول عمليات موحّد.
تدعم الآن: Excel/CSV, PDF (جداول), Word (.docx), نص (.txt).
الصور والملفات الممسوحة الفوضوية: تُوجَّه لطبقة الذكاء الاصطناعي (extract_with_ai)
  التي تُفعّل بمفتاح Claude API — نقطة تمديد واحدة، بدون تغيير باقي النظام.

الأعمدة الموحّدة: التاريخ، البيان، النوع، التصنيف، الطرف، المبلغ
"""
from __future__ import annotations
import os
import re
import logging
import pandas as pd

_log = logging.getLogger("almudir.extract")


def _load_dotenv():
    """يحمّل مفتاح .env تلقائياً — بلا مكتبات خارجية، ويتحمّل BOM وعلامات الاقتباس.
    لا يطغى على متغيرات البيئة الحقيقية إن وُجدت."""
    path = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env")
    if not os.path.exists(path):
        return
    try:
        with open(path, "r", encoding="utf-8-sig") as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith("#") or "=" not in line:
                    continue
                k, v = line.split("=", 1)
                k, v = k.strip(), v.strip().strip('"').strip("'")
                if k and k not in os.environ:
                    os.environ[k] = v
    except OSError:
        pass


_load_dotenv()

STD_COLS = ["التاريخ", "البيان", "النوع", "التصنيف", "الطرف", "المبلغ"]

# مرادفات الأعمدة (عربي + إنجليزي) — لمطابقة رؤوس الجداول المتنوعة
SYN = {
    "التاريخ": ["التاريخ", "تاريخ", "date", "day", "trans date", "transaction date", "posting date"],
    "البيان": ["البيان", "بيان", "الوصف", "وصف", "تفاصيل", "description", "details", "narration", "memo", "notes"],
    "النوع": ["النوع", "نوع", "type", "dr/cr", "debit/credit"],
    "التصنيف": ["التصنيف", "تصنيف", "البند", "الفئة", "category", "class", "account"],
    "الطرف": ["الطرف", "الجهة", "العميل", "المورد", "الاسم", "اسم الموظف", "party", "customer", "vendor", "name", "employee", "payee", "beneficiary"],
    "المبلغ": ["المبلغ", "مبلغ", "القيمة", "قيمة", "amount", "value", "total", "الراتب", "راتب", "salary", "أجر", "wage", "الفعلي", "actual", "المصروف"],
    # أعمدة مساعدة لاشتقاق النوع/المبلغ
    "_مدين": ["مدين", "debit", "withdrawal", "سحب", "منصرف"],
    "_دائن": ["دائن", "credit", "deposit", "إيداع", "وارد"],
}

INCOME_HINT = re.compile(r"دخل|إيراد|مبيع|إيداع|وارد|دائن|credit|deposit|income|sale", re.I)


class ExtractionError(Exception):
    pass


# ---------- كشف نوع الملف (يوجّه لتحليل مخصّص لكل نوع) ----------
def detect_ftype(columns) -> str:
    """يصنّف الملف من رؤوس الأعمدة: statement | payroll | budget.
    الافتراضي statement (يشمل كشوف الحساب وقوائم المصروفات — تحليل تدفق نقدي)."""
    cols = " ".join(str(c).strip().lower() for c in columns)
    def has(*ws): return any(w in cols for w in ws)
    # كشف رواتب: اسم/موظف + راتب/أجر
    if has("راتب", "salary", "أجر", "wage", "payroll") and \
       has("موظف", "employee", "staff", "الاسم", "اسم", "name"):
        return "payroll"
    # ميزانية/موازنة: مخطط + فعلي، أو كلمة موازنة/budget
    if (has("مخطط", "planned", "متوقع", "forecast", "الموازنة", "موازنة", "budget") and
            has("فعلي", "actual")) or has("موازنة", "budget"):
        return "budget"
    return "statement"


# ---------- المحلّل الحتمي للكشوف المنظّمة (راجحي وما شابهه) ----------
# نمط العملية: رصيد SAR | دائن SAR | مدين SAR | (تصنيف/وصف) | تاريخ YYYY/MM/DD
# حتمي بالكامل: أي عدد صفحات في أقل من ثانية، بلا ذكاء اصطناعي، بلا استهلاك ذاكرة، مجاناً.
_TX_BLOCK = re.compile(
    r"([\d,]+\.\d{2})\s*SAR\s+([\d,]+\.\d{2})\s*SAR\s+([\d,]+\.\d{2})\s*SAR"
    r"([^\n]*)\n(.*?)(\d{4}/\d{2}/\d{2})", re.DOTALL)

# قاموس مصطلحات البنوك السعودية → (تصنيف عربي نظيف، طرف نظيف، هل هو "مصدر عام" لا عميل؟)
# محلي بالكامل (بلا ذكاء اصطناعي خارجي) — حمايةً لخصوصية بيانات العميل (PDPL).
# الترتيب من الأخص للأعم. "المصدر العام" = إيداع نقدي/تحويل/رسوم — ليس عميلاً يُخشى فقدانه.
_TERM_RULES = [
    (r"راتب|رواتب|payroll|salary",                              "رواتب",           None,                    False),
    # الغرامات قبل «سداد الفواتير» — «غرامة تأخير سداد» كانت تذوب في الرسوم فتختفي
    (r"غرامة|مخالفة|penalty|late\s*fee",                        "غرامات ومخالفات", "غرامات ومخالفات",       True),
    # الأصول الرأسمالية (شراء لمرة واحدة — ليست مصروفاً تشغيلياً متكرراً)
    (r"شراء\s*(?:سيارة|مركبة|معدات|أثاث|عقار)|أصل\s*ثابت|أصول\s*ثابتة",
     "أصول (شراء لمرة)", "أصول (شراء لمرة)", True),
    (r"محطة|وقود|fuel|petrol|الدريس|ساسكو|aramco|نفط",           "وقود ومحروقات",   "محطات وقود",            True),
    (r"إيجار|ايجار|\brent\b|عقار",                               "إيجار",           "إيجار",                 True),
    (r"سناب|snap|قوقل|google|meta|facebook|tiktok|إعلان|تسويق|ads", "تسويق وإعلان",  None,                    False),
    (r"sadad|سداد|فواتير|فاتورة|كهرباء|اتصالات|\bstc\b|موبايلي|زين|مياه", "سداد فواتير", "سداد فواتير وخدمات", True),
    (r"sancbk|\bncbk\b|رسوم|عمول|\bfee\b|\bvat\b|ضريب|charge|خدمات بنك|مصاريف بنك", "رسوم وضرائب", "رسوم وخدمات بنكية", True),
    (r"نقاط البيع|point of sale|\bpos\b|mada|مدى|أثير|شراء عبر|purchase|pur\b", "مشتريات ونقاط بيع", "مشتريات ونقاط بيع", True),
    (r"toacct|fracct|تحويل|transfer|حوال|\bto:|\bfr:",           "تحويلات",         "تحويلات بنكية",         True),
    (r"cash deposit|إيداع|صراف|sanabil|نقدي|deposit",            "إيداع نقدي",      "إيداعات نقدية",         True),
]

# تسميات المصادر العامة — تُستبعَد من تحليل "تركّز العملاء" (ليست عملاء يمكن فقدانهم).
GENERIC_SOURCES = {"إيداعات نقدية", "تحويلات بنكية", "رسوم وخدمات بنكية",
                   "مشتريات ونقاط بيع", "سداد فواتير وخدمات", "محطات وقود", "إيجار",
                   "دخل غير مصنّف", "غير محدد", "سحوبات نقدية", "مصروفات أخرى",
                   "رواتب", "تسويق وإعلان", "وقود ومحروقات",
                   "عملاء متنوعون", "عملاء", "عميل", "متنوعون", "customers", "various"}

# ---------- طبقة "تشغيلي مقابل غير تشغيلي" (فخ الربح الجوهري) ----------
# رأس المال، التمويل، سحوبات المالك، سداد القروض، التحويلات الداخلية، وتحصيل الضريبة
# ليست دخلاً أو مصروفاً تشغيلياً — تُقلب المؤشرات (تركّز عملاء مخفّف، هامش كاذب) لو خُلطت.
# محلي بالكامل (regex)، يُطبَّق على أي بيانات مستخرجة بغضّ النظر عن مصدرها (AI أو حتمي).
NON_OPERATING = {"رأس مال", "تمويل بنكي", "سحب شخصي (المالك)", "سداد تمويل",
                 "تحويل داخلي بين الحسابات", "ضريبة محصّلة (تحصيل لا مصروف)",
                 "حساب الشريك (غير تشغيلي)", "شيك مرتجع (لا يُحتسب)",
                 "وارد غير مُوضّح (يحتاج توضيح)"}

_NON_OP_RULES = [
    (r"رأس\s*مال|capital\s*injection|ضخ\s*رأس|استثمار من الشريك|إيداع رأس مال", "رأس مال", True),
    # تمويل إسلامي/بنكي (مرابحة/تورّق/إجارة) = قرض داخل، ليس إيراداً تشغيلياً
    (r"مرابحة|تورّق|تورق|إجارة\s*منتهية|تمويل\s*بنكي|تمويل من|دفعة تمويل|"
     r"bank financing|loan disbursement|قرض\s*جديد", "تمويل بنكي", True),
    # حساب الشريك الجاري (ضخ/سحب من الشريك) — تمويل مالك لا مبيعات
    (r"حساب\s*الشريك\s*الجاري|الشريك\s*الجاري|جاري\s*الشريك|دفعة من حساب الشريك",
     "حساب الشريك (غير تشغيلي)", True),
    (r"سحب\s*شخصي|سحب\s*الشريك|owner draw|مسحوبات شخصية|مصروفات الشريك", "سحب شخصي (المالك)", False),
    (r"سداد\s*قسط\s*(?:تمويل|مرابحة)|سداد\s*قرض|قسط\s*تمويل|قسط\s*مرابحة|"
     r"loan\s*(?:installment|repayment)", "سداد تمويل", False),
    # محدَّد بدقة: تحويل المالك بين حساباته/حسابات مؤسسته الخاصة — لا "التحويل الداخلي"
    # كمصطلح بنكي عام (بعض البنوك تستخدمه لأي تحويل بين عملائها، وهذا دخل حقيقي).
    # نشترط ذكر «حساب المؤسسة/الشركة/آخر» صراحةً حتى لا نبتلع تحويلات العملاء الحقيقية.
    (r"(?:من|إلى|لـ|بين)\s*حساب(?:ات)?\s*(?:المؤسسة|الشركة|المالك|آخر|الآخر|الثاني)|"
     r"حساب\s*المؤسسة\s*الآخر|لحسابي\s*الآخر|"
     r"internal\s*transfer\s*between\s*(?:own|company)\s*accounts",
     "تحويل داخلي بين الحسابات", False),
    (r"ضريبة القيمة المضافة|zatca|هيئة الزكاة|vat remit", "ضريبة محصّلة (تحصيل لا مصروف)", False),
]


def tag_non_operating(df: pd.DataFrame) -> pd.DataFrame:
    """يعيد تصنيف الصفوف غير التشغيلية (رأس مال/تمويل/سحوبات/سداد قروض/تحويل داخلي/ضريبة)
    بغضّ النظر عن مصدر الاستخراج — طبقة موحّدة تُطبَّق على أي بيانات بعد الاستخراج."""
    if df.empty:
        return df
    # الرصيد الافتتاحي ليس عملية — نلتقط قيمته ونُسقطه، وإلا احتُسب «دخلاً» وهمياً
    # (بالضبط ما وقع فيه التقرير: رصيد افتتاحي 310,000 ظهر كإيراد 9%). ونحفظه لإعادة
    # بناء الرصيد الجاري (كشف السحب على المكشوف) لاحقاً.
    ob_mask = df["البيان"].astype(str).str.contains(
        r"رصيد\s*افتتاح|رصيد\s*سابق|رصيد\s*مُدوّر|opening\s*balance|balance\s*b/?f|\bb/f\b",
        regex=True, case=False, na=False)
    opening_balance = float(df.attrs.get("opening_balance", 0.0) or 0.0)  # لا نطمس قيمة قارئ سابق
    if ob_mask.any():
        opening_balance = opening_balance or float(df[ob_mask]["المبلغ"].iloc[0])
        attrs = dict(df.attrs)
        df = df[~ob_mask].copy()
        df.attrs.update(attrs)
    df.attrs["opening_balance"] = opening_balance
    blob = (df["البيان"].astype(str) + " " + df["الطرف"].astype(str)).str.lower()
    for pat, cat, _is_income in _NON_OP_RULES:
        mask = blob.str.contains(pat, regex=True, case=False, na=False)
        if mask.any():
            df.loc[mask, "التصنيف"] = cat
            df.loc[mask, "الطرف"] = cat
    # وارد غير مُوضّح: مرجع بنكي مبهم (IPS/TRF REF/QR/رموز) بلا اسم جهة حقيقي.
    # المبدأ: «نعترف لا نخمّن» — لا نحتسبه إيراداً تشغيلياً ولا نصنّفه «مبيعات» بثقة كاذبة،
    # بل نستبعده ونعرضه للمستخدم ليوضّحه. (خطأ واثق يقتل الثقة أسرع من الاعتراف بالجهل.)
    has_named_entity = blob.str.contains(r"شركة|مؤسسة|مجموعة", regex=True, case=False, na=False)
    has_merchant = df["الطرف"].apply(lambda p: _clean_merchant(str(p)) is not None)
    vague_pat = (r"\bips\b|تحويل\s*وارد|طرف\s*خارجي|غير\s*محدد|trf\s*ref|cr-?dr|qr\s*pymt|"
                 r"ips\s*cr|//\s*xx|3abc|[a-z]{2,}\s*ref\s*\d|ref\s*\d{3,}")
    is_vague_desc = blob.str.contains(vague_pat, regex=True, case=False, na=False)
    generic_party = df["الطرف"].astype(str).str.strip().isin(
        {"عميل", "تحويل", "طرف أخرى", "عملية مصرفية", "تاجر", "مورد", "غير محدد", ""})
    vague_in = (df["النوع"] == "دخل") & is_vague_desc & ~has_named_entity & (~has_merchant | generic_party)
    df.loc[vague_in, "التصنيف"] = "وارد غير مُوضّح (يحتاج توضيح)"
    df.loc[vague_in, "الطرف"] = "وارد غير مُوضّح (يحتاج توضيح)"

    # شيك مرتجع: إيداع شيك ثم ارتجاعه بنفس المبلغ والطرف = دخل وهمي (لا يُحتسب إيراداً).
    df = _tag_returned_cheques(df)
    return df


def _tag_returned_cheques(df: pd.DataFrame) -> pd.DataFrame:
    """يبطل أزواج (إيداع شيك ↔ ارتجاع شيك) بنفس الطرف والمبلغ — إيداعٌ لم يُصرف فعلاً.
    يمنع احتساب الـ65,000 المرتجعة كإيراد تشغيلي، ويهيّئها لكشف الشذوذ لاحقاً."""
    bayan = df["البيان"].astype(str)
    reversed_mask = bayan.str.contains(r"ارتجاع\s*شيك|شيك\s*مرتجع|شيك\s*مرتد|شيك\s*راجع|returned\s*che|bounced",
                                       regex=True, case=False, na=False)
    if not reversed_mask.any():
        return df
    for _, rev in df[reversed_mask].iterrows():
        party, amt = str(rev["الطرف"]), float(rev["المبلغ"])
        # الإيداع المقابل: نفس الطرف والمبلغ، دخل، ووصفه إيداع شيك
        dep = ((df["النوع"] == "دخل") & (df["الطرف"].astype(str) == party) &
               (df["المبلغ"] == amt) &
               df["البيان"].astype(str).str.contains(r"إيداع\s*شيك|شيك", regex=True, na=False))
        df.loc[dep | (df.index == rev.name), "التصنيف"] = "شيك مرتجع (لا يُحتسب)"
    return df


# رموز مصرفية ليست أسماء جهات — «TRF REF» ليس تاجراً (يمنع تخمين مرجع بنكي كاسم عميل)
_BANK_CODE_TOKENS = {"trf", "ref", "ips", "qr", "pymt", "cr", "dr", "acct",
                     "transfer", "payment", "pos", "atm", "chq", "sadad"}


def _clean_merchant(desc: str) -> str | None:
    """اسم تاجر/جهة لاتيني نظيف من الوصف (للمشتريات والتحويلات المسمّاة)."""
    if not desc:
        return None
    m = re.search(r"[A-Za-z][A-Za-z&'.\- ]{4,40}", desc)
    if not m:
        return None
    name = re.sub(r"\s+", " ", m.group(0)).strip(" .,-")
    skip = ("cash deposit", "online purchase", "the amount", "agmt", "toacct", "fracct")
    if len(name) < 4 or name.lower() in skip:
        return None
    # كل كلماته رموز مصرفية (TRF REF / IPS CR TRANSFER) → ليس اسم جهة
    tokens = [t for t in re.split(r"[^A-Za-z]+", name.lower()) if t]
    if tokens and all(t in _BANK_CODE_TOKENS for t in tokens):
        return None
    return name[:40]


def _classify(blob: str, desc: str, is_income: bool):
    """يُرجع (التصنيف، الطرف) بلغة عربية نظيفة — بلا أكواد بنكية خام."""
    low = blob.lower()
    for pat, cat, party, _generic in _TERM_RULES:
        if re.search(pat, low):
            # الصراف الآلي: إيداع إن كان وارداً، وسحب نقدي إن كان صادراً
            if cat == "إيداع نقدي" and not is_income:
                return "سحب نقدي", "سحوبات نقدية"
            if party is None:                          # رواتب/تسويق: نحاول استخراج اسم، وإلا التصنيف
                party = _clean_merchant(desc) or cat
            return cat, party
    # غير معروف: تصنيف عام + محاولة اسم تاجر
    if is_income:
        return "دخل غير مصنّف", (_clean_merchant(desc) or "دخل غير مصنّف")
    return "مصروفات أخرى", (_clean_merchant(desc) or "مصروفات أخرى")


def _read_structured_pdf(path: str) -> pd.DataFrame:
    """يقرأ كشوف البنوك المنظّمة (نمط راجحي) حتمياً وبسرعة عبر PyMuPDF."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ExtractionError("PyMuPDF غير مثبّت")
    import unicodedata
    rows = []
    doc = fitz.open(path)
    try:
        for pg in doc:
            # تطبيع NFKC: يحوّل أشكال العرض العربية (ﺷﺮﺍﺀ) إلى عربي عادي (شراء)
            # حتى تطابق قواعد التصنيف العربية. (بعض بنوك تخزّن النص مُشكّلاً مسبقاً.)
            text = unicodedata.normalize("NFKC", pg.get_text())
            for m in _TX_BLOCK.finditer(text):
                _bal, cr, db, tail, desc, date = m.groups()
                cr = float(cr.replace(",", "")); db = float(db.replace(",", ""))
                if cr <= 0 and db <= 0:
                    continue
                is_income = cr > 0
                label = (tail or "").strip()
                blob = f"{label} {desc or ''}"
                cat, party = _classify(blob, desc, is_income)
                rows.append({
                    "التاريخ": date.replace("/", "-"),
                    "البيان": re.sub(r"\s+", " ", blob).strip()[:120],
                    "النوع": "دخل" if is_income else "مصروف",
                    "التصنيف": cat,
                    "الطرف": party,
                    "المبلغ": cr if is_income else db,
                })
    finally:
        doc.close()
    if len(rows) < 3:
        raise ExtractionError("لم يُتعرّف على تنسيق كشف منظّم في هذا الـPDF.")
    out = pd.DataFrame(rows)[STD_COLS]
    out.attrs["ftype"] = "statement"
    return out


# ---------- قارئ كشوف Excel البنكية (رأس الجدول مدفون تحت بيانات تعريفية) ----------
_HDR_TOKENS = ("التاريخ", "البيان", "مدين", "دائن", "الرصيد", "المرجع",
               "date", "description", "debit", "credit", "balance")


def _num(v) -> float:
    s = re.sub(r"[^\d.\-]", "", str(v))
    try:
        return float(s) if s not in ("", "-", ".") else 0.0
    except ValueError:
        return 0.0


def _party_from_desc(desc: str, is_income: bool) -> str | None:
    """اسم الطرف من البيان: «شركة البناء المتحدة - دفعة عقد» → شركة البناء المتحدة،
    و«تحويل راتب - م. س. ع» → م. س. ع (اسم المستفيد، يلزم لكشف الرواتب لكل موظف)."""
    parts = [p.strip() for p in re.split(r"\s+-\s+|—", str(desc)) if p.strip()]
    if not parts:
        return None
    # رواتب/تحويلات مسمّاة: المستفيد هو الجزء الأخير («تحويل راتب - م. س. ع» → م. س. ع؛
    # «تحويل صادر - م ع ش» → م ع ش). يلزم لكشف الرواتب لكل موظف والمدفوعات المتصاعدة لجهة.
    if re.search(r"راتب|رواتب|salary|payroll|تحويل\s*(?:صادر|وارد)|حوالة", parts[0], re.I) and len(parts) > 1:
        return parts[-1][:40]
    for p in parts:                                   # جهة مسمّاة (شركة/مؤسسة/بنك…)
        if re.search(r"شركة|مؤسسة|مجموعة|مصرف|بنك|محطة|مورد", p):
            return p[:40]
    return _clean_merchant(desc)


def _read_bank_xlsx(path: str) -> pd.DataFrame:
    """يقرأ كشف حساب Excel بصيغته البنكية: بيانات تعريفية (اسم/آيبان/رصيد افتتاحي)
    ثم رأس جدول (التاريخ/البيان/مدين/دائن/الرصيد) ثم العمليات. حتمي بالكامل.
    يلتقط الرصيد الافتتاحي والختامي وأدنى رصيد — لكشف السحب على المكشوف بدقة."""
    raw = pd.read_excel(path, header=None)
    hdr_i = None
    for i in range(min(25, len(raw))):
        vals = [str(v).strip().lower() for v in raw.iloc[i].tolist()]
        hits = sum(1 for v in vals if any(t in v for t in [t.lower() for t in _HDR_TOKENS]))
        if hits >= 3:
            hdr_i = i
            break
    if hdr_i is None:
        raise ExtractionError("لم يُتعرَّف على رأس جدول كشف بنكي في الملف.")

    # الرصيد الافتتاحي من البيانات التعريفية فوق الرأس
    opening = 0.0
    for i in range(hdr_i):
        cells = [str(v) for v in raw.iloc[i].tolist()]
        for j, cell in enumerate(cells):
            if re.search(r"رصيد.*افتتاح|opening\s*balance", cell, re.I):
                for cand in cells[j:j + 3]:
                    n = _num(cand)
                    if n:
                        opening = n
                        break

    header = [str(v).strip() for v in raw.iloc[hdr_i].tolist()]
    body = raw.iloc[hdr_i + 1:].copy()
    body.columns = header

    def col(*tokens):
        for c in header:
            lc = str(c).lower()
            if any(t in lc or t in str(c) for t in tokens):
                return c
        return None

    c_date, c_desc = col("التاريخ", "date"), col("البيان", "description", "الوصف")
    c_db, c_cr = col("مدين", "debit", "سحب"), col("دائن", "credit", "إيداع")
    c_bal = col("الرصيد", "balance")
    if not (c_date and c_desc and (c_db or c_cr)):
        raise ExtractionError("أعمدة الكشف البنكي غير مكتملة (تاريخ/بيان/مدين/دائن).")

    body["_dt"] = pd.to_datetime(body[c_date], errors="coerce")
    body = body[body["_dt"].notna()]                 # يسقط «الإجمالي» وسطور الإخلاء تلقائياً
    if body.empty:
        raise ExtractionError("لم تُقرأ عمليات من كشف الحساب.")

    rows = []
    balances = []
    for _, r in body.iterrows():
        db = _num(r[c_db]) if c_db else 0.0
        cr = _num(r[c_cr]) if c_cr else 0.0
        if db <= 0 and cr <= 0:
            continue
        is_income = cr > db
        amount = cr if is_income else db
        desc = str(r[c_desc]).strip()
        cat, party = _classify(desc, desc, is_income)
        named = _party_from_desc(desc, is_income)
        if named:
            party = named
        rows.append({
            "التاريخ": r["_dt"].strftime("%Y-%m-%d"),
            "البيان": desc[:120],
            "النوع": "دخل" if is_income else "مصروف",
            "التصنيف": cat, "الطرف": party, "المبلغ": amount,
        })
        if c_bal is not None:
            b = _num(r[c_bal])
            if b or str(r[c_bal]).strip() not in ("", "nan"):
                balances.append(b)

    if len(rows) < 3:
        raise ExtractionError("لم تُقرأ عمليات كافية من كشف الحساب.")
    out = pd.DataFrame(rows)[STD_COLS]
    out.attrs["ftype"] = "statement"
    out.attrs["opening_balance"] = opening
    if balances:                                     # عمود الرصيد الحقيقي — أدق من إعادة البناء
        out.attrs["closing_balance"] = balances[-1]
        out.attrs["min_balance"] = min(balances)
    return out


# ---------- المُرسِل حسب نوع الملف ----------
def _has_ai():
    return bool(os.environ.get("ANTHROPIC_API_KEY"))


# حد صفحات الـPDF قابل للضبط بالبيئة: افتراضياً 60 (آمن لذاكرة الخطة المجانية 512MB).
# الأكبر يُوجَّه لتصدير Excel (خفيف على الذاكرة، فوري). ارفعه على استضافة أكبر عبر PDF_MAX_PAGES.
# حد آمن لذاكرة 512MB: أكبر من ذلك ينهار الخادم، فنعطي رسالة نظيفة فوراً بدل الانهيار.
# يُرفع على استضافة أكبر عبر PDF_MAX_PAGES، أو يُلغى عملياً عند إضافة المحلّل الحتمي للجداول.
_PDF_MAX_PAGES = int(os.environ.get("PDF_MAX_PAGES", "20"))
_PDF_DETERMINISTIC_MAX = 5  # PDF أكبر من هذا يذهب مباشرة لمسار النص (أسرع من استخراج الجداول)


def _pdf_page_count(path: str) -> int:
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            return len(pdf.pages)
    except Exception:
        return 0


def extract(path: str) -> pd.DataFrame:
    """يستخرج العمليات ثم يطبّق طبقة 'تشغيلي مقابل غير تشغيلي' على أي مصدر
    (AI أو حتمي) — طبقة موحّدة لا تعتمد على مسار الاستخراج."""
    df = _extract_inner(path)
    ftype = df.attrs.get("ftype", "statement")
    keep_attrs = dict(df.attrs)              # نحفظ الرصيد الافتتاحي/الختامي عبر المعالجة
    if ftype == "statement":                 # الرواتب/الميزانية لهما منطق مختلف مسبقاً
        df = tag_non_operating(df)
    merged = dict(keep_attrs); merged.update(df.attrs)
    df.attrs.update(merged)
    df.attrs["ftype"] = ftype
    return df


def _extract_inner(path: str) -> pd.DataFrame:
    ext = os.path.splitext(path)[1].lower()

    # الصور دائماً عبر الذكاء الاصطناعي
    if ext in (".png", ".jpg", ".jpeg", ".webp", ".heic"):
        return extract_with_ai(path)

    # PDF: المحلّل الحتمي المنظّم أولاً (راجحي وشبهه) — سريع، أي حجم، مجاناً، بلا ذاكرة ثقيلة
    if ext == ".pdf":
        try:
            return _read_structured_pdf(path)
        except ExtractionError:
            pass   # ليس بالتنسيق المنظّم → نكمل للمسارات الأخرى
        n = _pdf_page_count(path)
        if n > _PDF_MAX_PAGES:
            raise ExtractionError(
                f"كشفك كبير ({n} صفحة) وبتنسيق غير مُتعرَّف عليه. للنتيجة الأدق والأسرع، "
                "ارفع كشف آخر شهر أو ربع، أو صدّره Excel من تطبيق بنكك — يُعالَج فوراً مهما بلغ حجمه.")
        if n > _PDF_DETERMINISTIC_MAX and _has_ai():
            return extract_with_ai(path)   # كشف صغير غير منظّم → مسار الذكاء الاصطناعي

    # الصيغ المهيكلة: نجرّب القراءة الحتمية أولاً، ونرجع للذكاء الاصطناعي عند الفشل
    try:
        if ext in (".xlsx", ".xls"):
            try:
                return _normalize(pd.read_excel(path), source=ext)
            except ExtractionError:
                # كشف بنكي مُصدَّر (رأس الجدول مدفون تحت بيانات تعريفية) → قارئ حتمي مخصّص.
                # حتمي = نفس الأرقام في كل تشغيل، بلا تكلفة API، ويقرأ عمود الرصيد الحقيقي.
                return _read_bank_xlsx(path)
        if ext == ".csv":
            return _normalize(pd.read_csv(path), source=ext)
        if ext == ".pdf":
            return _normalize(_read_pdf_table(path), source=ext)
        if ext == ".docx":
            return _normalize(_read_docx_table(path), source=ext)
        if ext == ".txt":
            return _normalize(_read_txt_table(path), source=ext)
    except ExtractionError:
        if _has_ai():
            return extract_with_ai(path)   # كشف فوضوي / بلا جدول → الذكاء الاصطناعي
        raise
    raise ExtractionError(f"صيغة غير مدعومة: {ext}")


# ---------- التطبيع (مطابقة الأعمدة + اشتقاق النوع) ----------
def _match_col(cols, targets):
    norm = {str(c).strip().lower(): c for c in cols}
    for t in targets:
        for k, orig in norm.items():
            if t.lower() == k or t.lower() in k:
                return orig
    return None


def _normalize(df: pd.DataFrame, source: str) -> pd.DataFrame:
    if df is None or df.empty:
        raise ExtractionError("لم يُعثر على جدول عمليات صالح في الملف.")
    cols = list(df.columns)
    mapping = {}
    for std in STD_COLS + ["_مدين", "_دائن"]:
        found = _match_col(cols, SYN[std])
        if found is not None:
            mapping[std] = found

    if "المبلغ" not in mapping and not ("_مدين" in mapping or "_دائن" in mapping):
        raise ExtractionError(
            "ما قدرت أتعرف على عمود المبلغ. تأكد أن الملف فيه جدول واضح، "
            "أو استخدم القالب الجاهز (التاريخ، البيان، النوع، التصنيف، الطرف، المبلغ).")

    out = pd.DataFrame()
    out["التاريخ"] = df[mapping["التاريخ"]] if "التاريخ" in mapping else pd.NaT
    out["البيان"] = df[mapping["البيان"]] if "البيان" in mapping else ""
    out["التصنيف"] = df[mapping["التصنيف"]] if "التصنيف" in mapping else "غير مصنّف"
    out["الطرف"] = df[mapping["الطرف"]] if "الطرف" in mapping else "غير محدد"

    # المبلغ والنوع — عدة سيناريوهات
    if "المبلغ" in mapping:
        amt = pd.to_numeric(df[mapping["المبلغ"]].astype(str).str.replace(r"[^\d.\-]", "", regex=True),
                            errors="coerce").fillna(0.0)
        if "النوع" in mapping:
            out["النوع"] = df[mapping["النوع"]].astype(str)
            out["المبلغ"] = amt.abs()
        else:
            # اشتقاق النوع من إشارة المبلغ (سالب = مصروف)
            out["النوع"] = amt.apply(lambda v: "مصروف" if v < 0 else "دخل")
            out["المبلغ"] = amt.abs()
    else:
        # عمودا مدين/دائن منفصلان
        debit = pd.to_numeric(df[mapping["_مدين"]], errors="coerce").fillna(0.0) if "_مدين" in mapping else 0.0
        credit = pd.to_numeric(df[mapping["_دائن"]], errors="coerce").fillna(0.0) if "_دائن" in mapping else 0.0
        net = (credit if isinstance(credit, (int, float)) else credit) - \
              (debit if isinstance(debit, (int, float)) else debit)
        out["النوع"] = net.apply(lambda v: "دخل" if v >= 0 else "مصروف")
        out["المبلغ"] = net.abs()

    # لو النوع نصّي غامض، نطبّعه لدخل/مصروف عبر التلميحات
    out["النوع"] = out["النوع"].apply(
        lambda s: "دخل" if INCOME_HINT.search(str(s)) else ("مصروف" if str(s).strip() else "مصروف"))

    # نوع الملف يوجّه التحليل. الرواتب والميزانية = مصروفات/التزامات (تدفق خارج).
    ftype = detect_ftype(cols)
    if ftype == "payroll":
        out["النوع"] = "مصروف"
        if "التصنيف" not in mapping:
            out["التصنيف"] = "رواتب"
    elif ftype == "budget":
        out["النوع"] = "مصروف"

    out = out[out["المبلغ"] > 0].reset_index(drop=True)
    if out.empty:
        raise ExtractionError("الجدول لا يحتوي على مبالغ صالحة.")
    out = out[STD_COLS]
    out.attrs["ftype"] = ftype
    return out


# ---------- قرّاء الصيغ ----------
def _read_pdf_table(path):
    import pdfplumber
    rows = []
    header = None
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            for table in (page.extract_tables() or []):
                if not table:
                    continue
                if header is None:
                    header = [str(h).strip() if h else f"col{i}" for i, h in enumerate(table[0])]
                    body = table[1:]
                else:
                    body = table
                for r in body:
                    if r and any(c for c in r):
                        rows.append(r[:len(header)])
    if not header or not rows:
        raise ExtractionError("ما لقيت جدول في ملف PDF. لو الملف صورة/ممسوح، يحتاج طبقة الذكاء الاصطناعي.")
    return pd.DataFrame(rows, columns=header)


def _read_docx_table(path):
    import docx
    doc = docx.Document(path)
    if not doc.tables:
        raise ExtractionError("ملف Word لا يحتوي جدولاً. ضع البيانات في جدول بأعمدة واضحة.")
    t = doc.tables[0]
    header = [c.text.strip() for c in t.rows[0].cells]
    rows = [[c.text.strip() for c in row.cells] for row in t.rows[1:]]
    return pd.DataFrame(rows, columns=header)


def _read_txt_table(path):
    # نص مفصول بفواصل/تبويب
    for sep in [",", "\t", ";", "|"]:
        try:
            df = pd.read_csv(path, sep=sep)
            if df.shape[1] >= 3:
                return df
        except Exception:
            continue
    raise ExtractionError("ما قدرت أقرأ الملف النصي كجدول.")


# ---------- طبقة الذكاء الاصطناعي: قارئ كشوف البنوك الذكي ----------
# النموذج الافتراضي: Opus 4.8 — الأدق في قراءة الأرقام العربية من الكشوف الممسوحة.
# الدقة هنا هي المنتج (حصننا ضد ChatGPT)؛ التكلفة لكل مستند ضئيلة أمام الاشتراك الشهري.
# قابل للتبديل عبر ANTHROPIC_MODEL لمن يريد نموذجاً أرخص عن قصد.
# نموذج سريع للاستخراج المنظّم — Haiku 4.5 سريع ودقيق وأرخص بكثير من Opus.
# (Opus كان يجعل تحليل كشف حقيقي يستغرق دقيقة–دقيقتين ويبدو معلّقاً.)
AI_MODEL = os.environ.get("ANTHROPIC_MODEL", "claude-haiku-4-5")

_IMAGE_MEDIA = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
                ".webp": "image/webp"}

# مخطط الإخراج المنظّم — يضمن JSON صالح دائماً
_TX_SCHEMA = {
    "type": "object",
    "properties": {
        "transactions": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "date": {"type": "string", "description": "تاريخ العملية YYYY-MM-DD"},
                    "description": {"type": "string", "description": "بيان العملية"},
                    "type": {"type": "string", "enum": ["دخل", "مصروف"]},
                    "category": {"type": "string", "description": "تصنيف عربي: رواتب/إيجار/تسويق/مبيعات/مشتريات/تحويلات/رسوم بنكية/وقود/تشغيل/أخرى"},
                    "party": {"type": "string", "description": "اسم الطرف أو التاجر"},
                    "amount": {"type": "number", "description": "المبلغ موجب"},
                },
                "required": ["date", "description", "type", "category", "party", "amount"],
                "additionalProperties": False,
            },
        }
    },
    "required": ["transactions"],
    "additionalProperties": False,
}

_AI_PROMPT = (
    "أنت محلل مالي خبير في كشوف الحسابات البنكية السعودية (الراجحي، الأهلي، الإنماء، "
    "بنك ساب، الرياض) وأنظمة المحاسبة (قيود، دفترة). استخرج كل العمليات من هذا المستند.\n"
    "لكل عملية: حدّد النوع (دخل أو مصروف)، واستنتج تصنيفاً عربياً مناسباً من البيان أو اسم "
    "التاجر (مثال: 'محطة الدريس' → وقود، 'راتب' → رواتب، 'إيجار' → إيجار، تحويل من عميل → "
    "مبيعات)، واسم الطرف/التاجر، والمبلغ كرقم موجب، والتاريخ بصيغة YYYY-MM-DD.\n"
    "تجاهل عمود الرصيد الجاري إن وُجد — نريد مبلغ العملية فقط. لا تخترع عمليات غير موجودة. "
    "أعِد كل العمليات التي تراها."
)


def _pdf_all_text(path: str) -> str:
    """يستخرج نص كل صفحات الـPDF (للكشوف النصّية — الأغلبية)."""
    import pdfplumber
    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            if t.strip():
                parts.append(t)
    return "\n".join(parts)


# القيد الحقيقي على حجم الجزء = مخرجات النموذج (≤16 ألف توكن).
# عملية واحدة ≈ ~45 توكن مخرجات، فنحدّ كل جزء بعدد أسطر آمن (+حدّ أحرف للأسطر الطويلة).
_CHUNK_LINES = 220
_CHUNK_CHARS = 45000
_MAX_CHUNKS = 14  # يكفي ~3000 عملية (مئات الصفحات)؛ أبعد منه نوجّه لتصدير Excel


def _chunk_text(text: str, max_lines: int = _CHUNK_LINES, max_chars: int = _CHUNK_CHARS):
    """يقسّم النص على حدود الأسطر — بحدّ أقصى للأسطر (يضبط المخرجات) وللأحرف (يضبط المدخلات)."""
    lines = text.splitlines(keepends=True)
    chunks, cur, n_lines, n_chars = [], [], 0, 0
    for ln in lines:
        if cur and (n_lines >= max_lines or n_chars + len(ln) > max_chars):
            chunks.append("".join(cur)); cur, n_lines, n_chars = [], 0, 0
        cur.append(ln); n_lines += 1; n_chars += len(ln)
    if cur:
        chunks.append("".join(cur))
    return chunks or [text]


def _call_model(client, anthropic, content_blocks) -> list:
    """مكالمة واحدة للنموذج → قائمة عمليات خام. تعالج الأخطاء وحالات التوقف."""
    import json
    try:
        resp = client.messages.create(
            model=AI_MODEL, max_tokens=16000, system=_AI_PROMPT,
            messages=[{"role": "user", "content": content_blocks}],
            output_config={"format": {"type": "json_schema", "schema": _TX_SCHEMA}},
        )
    except anthropic.AuthenticationError:
        _log.error("Anthropic auth failed — invalid API key")
        raise ExtractionError("الخدمة غير متاحة مؤقتاً. حاول بعد قليل.")
    except anthropic.APITimeoutError:
        raise ExtractionError("المعالجة أخذت وقتاً أطول من المتوقع. جرّب تصدير الكشف كملف Excel — يُعالَج فوراً.")
    except anthropic.RateLimitError:
        raise ExtractionError("الخدمة مشغولة حالياً. أعد المحاولة بعد دقيقة.")
    except anthropic.BadRequestError as e:
        msg = str(getattr(e, "message", e))
        _log.error("Anthropic 400: %s", msg)
        if "too long" in msg or "maximum" in msg.lower() or "PDF pages" in msg:
            raise ExtractionError(
                "هذا الكشف ضخم. صدّره كملف Excel أو CSV من بنكك — يُعالَج فوراً وبدقة مهما كان حجمه.")
        raise ExtractionError("الخدمة غير متاحة مؤقتاً. حاول بعد قليل.")
    except anthropic.APIConnectionError:
        raise ExtractionError("تعذّر الاتصال بخدمة الذكاء الاصطناعي. تأكد من الإنترنت وحاول مرة أخرى.")
    except anthropic.APIError as e:
        _log.error("Anthropic API error: %s", getattr(e, "message", e))
        raise ExtractionError("الخدمة غير متاحة مؤقتاً. حاول بعد قليل.")

    if resp.stop_reason == "refusal":
        raise ExtractionError("تعذّرت معالجة هذا الملف. جرّب ملفاً آخر أو صيغة أوضح.")
    if resp.stop_reason == "max_tokens":
        raise ExtractionError(
            "هذا الكشف يحوي عمليات أكثر من أن تُعالَج دفعة. صدّره كملف Excel/CSV، أو ارفع فترة أقصر.")
    text = next((b.text for b in resp.content if b.type == "text"), "")
    try:
        return json.loads(text).get("transactions", [])
    except (json.JSONDecodeError, AttributeError):
        raise ExtractionError("ما قدرنا نقرأ العمليات من المستند. تأكد أنه كشف حساب واضح.")


def _extract_text_via_ai(client, anthropic, text: str) -> list:
    """يقسّم النص الطويل إلى أجزاء ويعالجها بالتوازي ثم يدمج — سريع ويتوسّع لأي حجم."""
    chunks = _chunk_text(text)
    if len(chunks) > _MAX_CHUNKS:
        raise ExtractionError(
            "الكشف ضخم (مئات الصفحات). صدّره كملف Excel أو CSV من تطبيق بنكك — "
            "يُعالَج فوراً وبدقة مهما كان عدد صفحاته.")
    if len(chunks) == 1:
        return _call_model(client, anthropic,
                           [{"type": "text", "text": "محتوى كشف الحساب:\n\n" + chunks[0]}])

    # أجزاء متعددة → توازٍ (كل جزء مكالمة I/O مستقلّة) لتقليص الزمن للثلث
    from concurrent.futures import ThreadPoolExecutor
    def _one(ch):
        return _call_model(client, anthropic,
                           [{"type": "text", "text": "محتوى كشف الحساب (جزء):\n\n" + ch}])
    # توازٍ محدود (2): يوازن بين السرعة وذاكرة الخطة المجانية (512MB) لتفادي OOM/502
    workers = int(os.environ.get("AI_WORKERS", "2"))
    rows = []
    with ThreadPoolExecutor(max_workers=min(len(chunks), workers)) as ex:
        for part in ex.map(_one, chunks):     # يحافظ على الترتيب ويرفع أي ExtractionError
            rows += part
    return rows


def extract_with_ai(path: str) -> pd.DataFrame:
    """
    تحوّل أي ملف (صورة، PDF، كشف بنكي، مستند فوضوي) إلى جدول عمليات موحّد.
    الكشوف النصّية الطويلة تُقسَّم تلقائياً لتتجاوز حد سياق النموذج (يتوسّع لأي حجم).
    """
    if not os.environ.get("ANTHROPIC_API_KEY"):
        raise ExtractionError(
            "قراءة الصور والكشوف الممسوحة تحتاج تفعيل طبقة الذكاء الاصطناعي "
            "(مفتاح Claude API في متغير البيئة ANTHROPIC_API_KEY). "
            "مبدئياً استخدم Excel / CSV / PDF / Word فيه جدول واضح.")
    try:
        import anthropic
    except ImportError:
        raise ExtractionError("مكتبة anthropic غير مثبّتة: pip install anthropic")

    import base64
    ext = os.path.splitext(path)[1].lower()
    client = anthropic.Anthropic(timeout=100.0, max_retries=1)   # يفشل بلطف بدل التعليق للأبد

    if ext in _IMAGE_MEDIA:                                    # صورة كشف
        with open(path, "rb") as f:
            data = base64.standard_b64encode(f.read()).decode()
        rows = _call_model(client, anthropic, [{"type": "image", "source": {
            "type": "base64", "media_type": _IMAGE_MEDIA[ext], "data": data}}])
    elif ext == ".pdf":
        text = ""
        try:
            text = _pdf_all_text(path)
        except Exception:
            text = ""
        if len(text.strip()) >= 200:                          # PDF نصّي → نص مقسّم (بلا حد صفحات)
            rows = _extract_text_via_ai(client, anthropic, text)
        else:                                                 # PDF ممسوح (صورة) → مستند
            with open(path, "rb") as f:
                data = base64.standard_b64encode(f.read()).decode()
            rows = _call_model(client, anthropic, [{"type": "document", "source": {
                "type": "base64", "media_type": "application/pdf", "data": data}}])
    else:                                                     # xlsx/csv/txt فوضوي → نص مقسّم
        if ext in (".xlsx", ".xls"):
            text = pd.read_excel(path).to_csv(index=False)
        else:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
        rows = _extract_text_via_ai(client, anthropic, text)

    if not rows:
        raise ExtractionError("ما لقينا عمليات في المستند.")
    df = pd.DataFrame(rows).rename(columns={
        "date": "التاريخ", "description": "البيان", "type": "النوع",
        "category": "التصنيف", "party": "الطرف", "amount": "المبلغ"})
    return _normalize(df, source="ai")
